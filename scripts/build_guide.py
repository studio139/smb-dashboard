# -*- coding: utf-8 -*-
"""Build מדריך-שימוש.docx — a short Hebrew RTL user guide, deterministically.

Hebrew Word correctness is set at the XML level (the step most tools miss):
  * every Hebrew paragraph -> <w:bidi/> + right alignment
  * every Hebrew run       -> <w:rtl/> + Arial (ascii/hAnsi/cs)
  * RTL tables             -> <w:bidiVisual/>
  * Latin/number tokens inside RTL text are wrapped in LRM so they don't reorder;
    pure command/path lines are rendered as their own LTR paragraphs.

Run:  C:\\Python314\\python.exe scripts\\build_guide.py
"""
import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "מדריך-שימוש.docx")

FONT = "Arial"
LRI, PDI = "⁦", "⁩"  # LTR isolate open/close — keeps Latin/paths LTR in RTL


def _ltr(s):
    """Wrap a Latin/number/path token in a balanced LTR isolate."""
    return LRI + str(s) + PDI
INK = RGBColor(0x1A, 0x1A, 0x1A)
BLUE = RGBColor(0x1F, 0x3A, 0x4D)   # PRIMARY accent (matches style_config.py)
MUTED = RGBColor(0x5C, 0x66, 0x70)
CODEBG = "EEF1F6"


def _set_font(run, size, bold, color, rtl):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = FONT
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT)
    if rtl:
        rpr.append(OxmlElement("w:rtl"))


def _bidi(p):
    ppr = p._p.get_or_add_pPr()
    ppr.append(OxmlElement("w:bidi"))


def para(doc, runs, align="right", rtl=True, space_after=6, space_before=0):
    """runs = list of (text, opts). opts: size, bold, color."""
    p = doc.add_paragraph()
    if rtl:
        _bidi(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else \
        (WD_ALIGN_PARAGRAPH.LEFT if align == "left" else WD_ALIGN_PARAGRAPH.CENTER)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    for text, opts in runs:
        r = p.add_run(text)
        _set_font(r, opts.get("size", 11), opts.get("bold", False),
                  opts.get("color", INK), rtl)
    return p


def heading(doc, text):
    return para(doc, [(text, {"size": 14, "bold": True, "color": BLUE})],
                space_before=12, space_after=6)


def bullet(doc, segments):
    """segments = list of (text, is_code) -> one RTL bullet with LRM-wrapped code."""
    runs = [("•  ", {"size": 11, "bold": True, "color": BLUE})]
    for text, is_code in segments:
        if is_code:
            runs.append((_ltr(text), {"size": 10.5, "bold": True, "color": BLUE}))
        else:
            runs.append((text, {"size": 11}))
    return para(doc, runs, space_after=4)


def code_line(doc, text):
    """A pure command line, rendered LTR so the path/command reads correctly."""
    p = para(doc, [(text, {"size": 10.5, "bold": True, "color": BLUE})],
             align="left", rtl=False, space_after=3)
    return p


def build():
    doc = docx.Document()
    # whole-section RTL at the XML level (python-docx Section has no RTL property)
    doc.sections[0]._sectPr.append(OxmlElement("w:bidi"))

    para(doc, [("מדריך שימוש — דשבורד פיננסי", {"size": 20, "bold": True, "color": BLUE})],
         align="right", space_after=2)
    para(doc, [("סטודיו מרים בסקון", {"size": 12, "color": MUTED})], space_after=10)
    para(doc, [("המנוע מחשב כל מספר ובונה את הדוח מחדש בכל הרצה — מראה קבוע, גמיש לנתונים. "
                "כל המספרים מהקוד; שום נתון אינו מומצא.", {"size": 11})], space_after=10)

    # 1
    heading(doc, "1. עבודה חודשית — היכן שמים את הקבצים")
    bullet(doc, [("בכל חודש מייצאים מ-", False), ("Fireberry", True),
                 (" את חמשת הקבצים ושמים אותם בתיקיית החודש, באותם שמות בדיוק:", False)])
    code_line(doc, r"inputs\2026-05\  →  לקוחות.xlsx · פרויקטים.xlsx · הכנסות.xlsx · שתפים.xlsx · משימות.xlsx")
    bullet(doc, [("הקובץ ", False), ("הוצאות קבועות.xlsx", True),
                 (" נשאר בשורש התיקייה ומתעדכן במקום (לא מעתיקים אותו לתיקיית החודש).", False)])
    bullet(doc, [("תיקיית החודש היא גם הארכיון — הקבצים נשארים בה, אין שלב גיבוי נפרד.", False)])

    # 2
    heading(doc, "2. הרצת דוח — חודשי / רבעוני / שנתי")
    bullet(doc, [("פותחים טרמינל בתיקייה ומריצים אחת מהפקודות:", False)])
    code_line(doc, r"C:\Python314\python.exe scripts\run.py 2026-05      (חודשי)")
    code_line(doc, r"C:\Python314\python.exe scripts\run.py 2026-Q1      (רבעוני)")
    code_line(doc, r"C:\Python314\python.exe scripts\run.py 2026         (שנתי)")
    bullet(doc, [("הפלט נכתב לתיקייה נפרדת לכל תקופה תחת ", False), ("outputs\\", True),
                 (" — למשל ", False), ("outputs\\2026-04\\", True),
                 (": קובץ ", False), (".xlsx", True), (" ותצוגת ", False), (".html", True),
                 (" בלבד (ללא PDF). שום קובץ אינו נשאר ישירות תחת ", False), ("outputs\\", True), (".", False)])
    bullet(doc, [("ה-", False), ("html", True),
                 (" הוא דשבורד אינטראקטיבי עם טאבים (סקירה · כספים · פניות והמרה · משימות · פירוט) "
                  "שנפתח בכל דפדפן — גם בלי אינטרנט; כל המספרים זהים לאקסל.", False)])
    bullet(doc, [("בסיום הריצה הדוח מתפרסם אוטומטית ל-", False), ("Drive", True),
                 (" והפרויקט מגובה ל-", False), ("GitHub", True), (" — פירוט בסעיף 6.", False)])

    # 3
    heading(doc, "3. ישיבת הרבעון וטבלת היעדים")
    bullet(doc, [("בסוף כל רבעון נוצרת אוטומטית תבנית ", False), ("Word", True),
                 (" ריקה בתיקיית חודש הסגירה (למשל ", False), ("inputs\\2026-03\\", True), (").", False)])
    bullet(doc, [("ממלאים את ", False), ("טבלת היעדים", False),
                 (" — מספר לכל מדד (פניות · עסקאות · המרה · מחזור · רווח · תקרת הוצאות). "
                  "שמות המדדים מתחברים אוטומטית לדוח.", False)])
    bullet(doc, [("התובנות והיוזמות הן טקסט חופשי — לא נשלף מהן מספר.", False)])
    bullet(doc, [("מהדוח הבא ואילך מוצג בלוק ", False), ("עמידה ביעדים", False),
                 (" (בפועל מול יעד). היעדים נשמרים עד שממלאים ", False), ("Word", True), (" חדש.", False)])

    # 4
    heading(doc, "4. איך קוראים את הדוח")
    bullet(doc, [("שתי תמונות הכנסה, נפרדות לעולם: ", False), ("פייפליין", False),
                 (" — כסף שנסגר (שכ\"ט מפרויקטים ושת\"פים), מול ", False), ("תזרים", False),
                 (" — כסף שנכנס (מטבלת ההכנסות).", False)])
    bullet(doc, [("מע\"מ: הכנסה נטו = ברוטו ÷ 1.18 רק כשחשבונית = כן; ריק או “-” = ללא חשבונית.", False)])
    bullet(doc, [("המרת קוהורט — לפי חודש הגעת הליד; ~3 החודשים האחרונים מסומנים ", False),
                 ("“עוד מבשיל”", False),
                 (". הכותרת מול היעד היא ", False), ("שיעור המרה מיוצב", False),
                 (" (קוהורטות בשלות בלבד).", False)])
    bullet(doc, [("זמן המרה — חציון הימים מהגעת הליד ועד תחילת העסקה ", False),
                 ("(מבוסס על עסקאות שחוברו לליד; מוטה כלפי מטה עד שהיסטוריית הלידים תתרחב).", False)])
    bullet(doc, [("כל בלוק אנליטי מוצג חודש-אחר-חודש + עמודת סיכום (כמו שתי תמונות ההכנסה): "
                  "הוצאות ורווח, מקור הפניות, ומבצעי המשימות — כך רואים רווח והוצאות לכל חודש, לא רק לרבעון. "
                  "הרווח צבוע ירוק/אדום לפי סימן.", False)])
    bullet(doc, [("גיליון ", False), ("הוצאות", False),
                 (" מפרט כל שורת הוצאה לכל חודש בתקופה; סכום השורות = סך ההוצאות בכותרת.", False)])

    # 5
    heading(doc, "5. שיוך עסקאות ללידים — מדד אינפורמטיבי")
    bullet(doc, [("הבלוק מראה כמה מהעסקאות שנסגרו מקורן ב", False), ("ליד מתויג", False),
                 (" (חיבור לפי שם הלקוח).", False)])
    bullet(doc, [("היתר — הפניה, לקוח חוזר, פנייה ישירה, או לקוח שהגיע לפני חלון המעקב — ", False),
                 ("אינו שגיאה ואינו חוסר נתונים", False),
                 (". לכל פרויקט תמיד יש לקוח; זהו מדד הקשר בלבד.", False)])

    # 6
    heading(doc, "6. פרסום אוטומטי ל-Drive וגיבוי ל-GitHub")
    bullet(doc, [("בסוף כל ריצה, אחרי יצירת הדוח, מתבצעים אוטומטית שני צעדים — שניהם ", False),
                 ("רכים", False),
                 (": אם משהו נכשל (אין רשת או הזדהות) מודפסת הודעה והריצה ממשיכה, הדוח לעולם לא נחסם.", False)])
    bullet(doc, [("העלאה ל-", False), ("Drive", True),
                 (" — קובצי התקופה (", False), ("xlsx", True), (" ו-", False), ("html", True),
                 (") עולים ל-Drive המשותף, לתיקייה «דשבורדים» (תיקיית-משנה לכל תקופה) תחת "
                  "«ניתוחים וישיבות הנהלה». מתחבר עם החשבון ", False), ("studio@smb-arch.com", True), (".", False)])
    bullet(doc, [("אם ההעלאה מדלגת עם בקשת הזדהות — מריצים פעם אחת ", False), ("gws auth login", True),
                 (". אפליקציית ה-", False), ("OAuth", True),
                 (" כעת ב-Production, כך שההזדהות כבר לא פגה כל שבוע.", False)])
    bullet(doc, [("גיבוי ל-", False), ("GitHub", True),
                 (" — כל הפרויקט (קוד · קלטים · הוצאות · מדריך) מגובה אוטומטית (", False),
                 ("commit + push", True), (") ל-repo הפרטי ", False), ("studio139/smb-dashboard", True),
                 (". הפלט (", False), ("outputs", True),
                 (") אינו נשמר ב-git — הוא נוצר מחדש בכל ריצה וכבר נמצא ב-Drive.", False)])
    bullet(doc, [("בקיצור: כל ריצה מפרסמת את הדוח ומגבה את כל הפרויקט — בלי לגעת בכלום.", False)])

    para(doc, [("נבנה אוטומטית על ידי ", {"size": 9, "color": MUTED}),
               (_ltr("scripts/build_guide.py"), {"size": 9, "color": MUTED}),
               (" — לעריכה, שנים את הסקריפט ולא את הקובץ.", {"size": 9, "color": MUTED})],
         space_before=14)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    p = build()
    print("wrote " + os.path.relpath(p, ROOT))
